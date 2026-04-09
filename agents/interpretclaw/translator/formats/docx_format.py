"""DOCX format handler"""

from pathlib import Path
from .base import DocumentFormat

class DocxFormat(DocumentFormat):
    """Handler for .docx files"""
    
    @property
    def supported_extensions(self) -> list:
        return ['.docx']
    
    def extract(self, file_path: Path) -> str:
        """Extract text from DOCX file"""
        try:
            import docx
            doc = docx.Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return '\n'.join(paragraphs)
        except ImportError:
            raise Exception("python-docx not installed. Run: pip install python-docx")
        except Exception as e:
            raise Exception(f"Failed to read DOCX: {e}")
    
    def save(self, file_path: Path, content: str, target_lang: str) -> Path:
        """Save translated DOCX file"""
        try:
            import docx
            doc = docx.Document()
            for paragraph in content.split('\n'):
                if paragraph.strip():
                    doc.add_paragraph(paragraph)
            
            output_path = file_path.parent / f"{file_path.stem}_{target_lang}{file_path.suffix}"
            doc.save(str(output_path))
            return output_path
        except ImportError:
            output_path = file_path.parent / f"{file_path.stem}_{target_lang}.txt"
            output_path.write_text(content, encoding='utf-8')
            return output_path
