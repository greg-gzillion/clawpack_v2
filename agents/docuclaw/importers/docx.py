"""Microsoft Word (.docx) importer"""
from pathlib import Path
from importers.base import BaseImporter

class DocxImporter(BaseImporter):
    name = "docx"
    extensions = [".docx"]
    
    def import_file(self, file_path):
        p = Path(file_path)
        try:
            # Try to use python-docx if available
            from docx import Document
            doc = Document(p)
            text = "\n".join([para.text for para in doc.paragraphs])
            return {
                "text": text,
                "format": "docx",
                "paragraphs": len(doc.paragraphs)
            }
        except ImportError:
            return {
                "text": f"DOCX file: {p.name}\nInstall python-docx for full support: pip install python-docx",
                "format": "docx",
                "warning": "python-docx not installed"
            }
        except Exception as e:
            return {"text": f"Error reading DOCX: {e}", "format": "docx"}
