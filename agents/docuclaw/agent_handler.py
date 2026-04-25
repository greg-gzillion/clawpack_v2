"""A2A Handler for DocuClaw - Document Generator with Multi-Format Export"""
import sys, os
from pathlib import Path
from datetime import datetime

DOCUCLAW_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = DOCUCLAW_DIR.parent.parent
LLMCLAW_DIR = PROJECT_ROOT / "agents" / "llmclaw"
EXPORTS = PROJECT_ROOT / "exports"
sys.path.insert(0, str(PROJECT_ROOT))
sys.path.insert(0, str(DOCUCLAW_DIR))
sys.path.insert(0, str(LLMCLAW_DIR))

from shared.base_agent import BaseAgent
from commands.llm_enhanced import run as llm_run

class DocuClawAgent(BaseAgent):
    def __init__(self):
        super().__init__("docuclaw")

    def _call_llm(self, prompt, context=""):
        if context:
            prompt = "Reference context:\n" + context[:1500] + "\n\n" + prompt
        result = llm_run(prompt)
        return result if result and not result.startswith("Error:") else "Document generation failed"

    def _save(self, content, fmt, name):
        EXPORTS.mkdir(exist_ok=True)
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        fn = f"{name or 'doc'}_{ts}.{fmt}"
        path = EXPORTS / fn
        try:
            if fmt == "pdf":
                from fpdf import FPDF
                pdf = FPDF(); pdf.add_page(); pdf.set_font("Arial", size=12)
                for line in content.split(chr(10))[:200]: pdf.cell(200, 10, txt=line[:100], ln=True)
                pdf.output(str(path))
            elif fmt == "docx":
                from docx import Document
                doc = Document(); doc.add_heading(name or "Document", 0)
                for line in content.split(chr(10))[:500]:
                    if line.strip(): doc.add_paragraph(line[:500])
                doc.save(str(path))
            elif fmt == "html":
                html = f"<html><head><meta charset='utf-8'><title>{name or 'Document'}</title><style>body{{font-family:Arial;max-width:800px;margin:40px auto;padding:20px}}pre{{background:#f5f5f5;padding:15px;border-radius:5px;white-space:pre-wrap}}</style></head><body><h1>{name or 'DocuClaw Export'}</h1><pre>{content}</pre></body></html>"
                path.write_text(html, encoding="utf-8")
            else:
                path.write_text(content, encoding="utf-8")
            os.startfile(str(path))
            return fn
        except:
            path.write_text(content, encoding="utf-8")
            return fn

    def handle(self, task):
        self.track_interaction()
        task = task.strip()
        parts = task.split(maxsplit=1)
        cmd = parts[0].lower() if parts else ""
        args = parts[1] if len(parts) > 1 else ""
        query = args if args else task
        try:
            ctx = ""
            try: ctx = self.search_web(query, max_results=3)
            except: pass

            if cmd in ("/create", "/letter", "/report", "/memo", "/resume", "/proposal") and query:
                doc_type = cmd.replace("/", "")
                result = self._call_llm(f"Create a professional {doc_type} in Markdown format for: {query}", ctx)
                fmt = args.split()[-1] if args.split()[-1] in ("pdf","docx","html","md","txt") else "md"
                fn = self._save(result, fmt, doc_type)
                result = f"Saved: {fn}\n\n{result[:800]}"
            elif cmd == "/export" and args:
                parts2 = args.split(maxsplit=1)
                fmt = parts2[0]
                content = parts2[1] if len(parts2) > 1 else ""
                if content:
                    fn = self._save(content, fmt, "export")
                    result = f"Exported: {fn}"
                else:
                    result = "Usage: /export <pdf|docx|html|md> <content>"
            elif cmd == "/help":
                result = "DocuClaw - Document Generator\n  /create /letter /report /memo /resume /proposal <topic>\n  /export <pdf|docx|html|md> <content>\n  Auto-save: MD/PDF/DOCX/HTML"
            elif cmd == "/stats":
                result = f"DocuClaw | Multi-format | WebClaw | Interactions: {self.state.get('interactions', 0)}"
            else:
                result = self._call_llm(f"Create a document: {query}", ctx)
                fn = self._save(result, "md", "document")
                result = f"Saved: {fn}\n\n{result[:800]}"
            return {"status": "success", "result": str(result)}
        except Exception as e:
            return {"status": "error", "result": str(e)}

_agent = DocuClawAgent()
def process_task(task, agent=None):
    return _agent.handle(task)
